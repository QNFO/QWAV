"""Apply business-docs.css styling to pandoc-generated DOCX."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
import os

SRC = r"G:\My Drive\projects\QWAV\briefings\Richard Goodman - apoth3osis Meet-and-Greet.docx"
doc = Document(SRC)

# ── Colours ──
DARK   = RGBColor(0x1A, 0x1A, 0x1A)
MED    = RGBColor(0x33, 0x33, 0x33)
GRAY   = RGBColor(0x4D, 0x4D, 0x4D)
LGRAY  = RGBColor(0x66, 0x66, 0x66)
BORDER = RGBColor(0xCC, 0xCC, 0xCC)
THEAD  = RGBColor(0xF0, 0xF0, 0xF0)
TALT   = RGBColor(0xFA, 0xFA, 0xFA)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

FONT_NAME = "Segoe UI"
MONO_NAME = "Consolas"

# ── Page setup: A4 + margins ──
for section in doc.sections:
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin    = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(0.8)

# ── Helper: add bottom border to paragraph ──
def add_bottom_border(para, color=BORDER, sz=4):
    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '{:02X}{:02X}{:02X}'.format(*color))
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_run(run, size_pt=10.5, bold=False, color=DARK, italic=False, name=FONT_NAME):
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)

def set_spacing(para, after_pt=6, line_spacing=1.55):
    pf = para.paragraph_format
    pf.space_after = Pt(after_pt)
    pf.space_before = Pt(0)
    pf.line_spacing = line_spacing

# ── Style Heading 1 (top-level: 24pt / bold / black 3px bottom border) ──
h1_style = doc.styles['Heading 1']
h1_font = h1_style.font
h1_font.size = Pt(24)
h1_font.bold = True
h1_font.color.rgb = DARK
h1_font.name = FONT_NAME
h1_pf = h1_style.paragraph_format
h1_pf.space_before = Pt(42)
h1_pf.space_after = Pt(6)
h1_pf.line_spacing = 1.1

for p in doc.paragraphs:
    if p.style.name == 'Heading 1':
        add_bottom_border(p, DARK, 12)
        for run in p.runs:
            set_run(run, 24, True, DARK)

# ── Style Heading 2 (15pt / bold / 1px #ccc bottom border) ──
h2_style = doc.styles['Heading 2']
h2_style.font.size = Pt(15)
h2_style.font.bold = True
h2_style.font.color.rgb = DARK
h2_style.font.name = FONT_NAME
h2_style.paragraph_format.space_before = Pt(18)
h2_style.paragraph_format.space_after = Pt(4)
h2_style.paragraph_format.line_spacing = 1.2

for p in doc.paragraphs:
    if p.style.name == 'Heading 2':
        add_bottom_border(p, BORDER, 4)
        for run in p.runs:
            set_run(run, 15, True, DARK)

# ── Style Heading 3 (13pt / bold / #333) ──
h3_style = doc.styles['Heading 3']
h3_style.font.size = Pt(13)
h3_style.font.bold = True
h3_style.font.color.rgb = MED
h3_style.font.name = FONT_NAME
h3_style.paragraph_format.space_before = Pt(14)
h3_style.paragraph_format.space_after = Pt(4)

for p in doc.paragraphs:
    if p.style.name == 'Heading 3':
        for run in p.runs:
            set_run(run, 13, True, MED)

# ── Style Heading 4 (12pt / bold / #4d4d4d) ──
h4_style = doc.styles['Heading 4']
h4_style.font.size = Pt(12)
h4_style.font.bold = True
h4_style.font.color.rgb = GRAY
h4_style.font.name = FONT_NAME

for p in doc.paragraphs:
    if p.style.name == 'Heading 4':
        for run in p.runs:
            set_run(run, 12, True, GRAY)

# ── Body text: Segoe UI 10.5pt, 1.6 spacing, #1a1a1a ──
for p in doc.paragraphs:
    if p.style.name not in ('Heading 1','Heading 2','Heading 3','Heading 4','Title'):
        set_spacing(p, after_pt=6, line_spacing=1.55)
        for run in p.runs:
            if run.font.size is None or run.font.size < Pt(11):
                set_run(run, 10.5, False, DARK)

# ── Tables: alternating rows, header bg, bold first col, borders ──
def style_table(tbl):
    # Set borders on whole table
    tbl_pr = tbl._element.get_or_add_tblPr()
    tbl_borders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), 'D0D0D0')
        tbl_borders.append(el)
    tbl_pr.append(tbl_borders)

    for ri, row in enumerate(tbl.rows):
        for ci, cell in enumerate(row.cells):
            # Header row
            if ri == 0:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'F0F0F0')
                shading.set(qn('w:val'), 'clear')
                cell._element.get_or_add_tcPr().append(shading)
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_run(run, 10.5, True, DARK)
            else:
                # Alternating row color
                if ri % 2 == 0:
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), 'FAFAFA')
                    shading.set(qn('w:val'), 'clear')
                    cell._element.get_or_add_tcPr().append(shading)
                # Bold first column
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_run(run, 10.5, ci == 0, DARK)
            # Cell padding
            tc_pr = cell._element.get_or_add_tcPr()
            tc_mar = OxmlElement('w:tcMar')
            for side in ('top','left','bottom','right'):
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:w'), '60')
                el.set(qn('w:type'), 'dxa')
                tc_mar.append(el)
            tc_pr.append(tc_mar)

for tbl in doc.tables:
    style_table(tbl)

# ── Save ──
doc.save(SRC)
print(f"Done — {os.path.getsize(SRC)} bytes")
