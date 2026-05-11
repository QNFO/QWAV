"""Convert the polished meeting agenda from Markdown to DOCX and PDF (via LaTeX)."""
import re, os, subprocess

SRC = r"G:\My Drive\projects\QWAV\briefings\Richard Goodman - Meeting Agenda (Shared).md"
OUT_DIR = os.path.dirname(SRC)
BASE = os.path.join(OUT_DIR, "Richard Goodman - Meeting Agenda")
TEXLIVE_BIN = r"C:\texlive\2025\bin\windows"

def read_md(path):
    with open(path, 'r', encoding='utf-8') as f:
        return f.read()

# ── markdown block parser ────────────────────────────────────────────
def parse_md(text):
    lines = text.split('\n')
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.strip() == '':
            i += 1
            continue
        # Horizontal rule
        if re.match(r'^[-*_]{3,}\s*$', line.strip()):
            blocks.append({'type': 'hr'})
            i += 1
            continue
        # Fenced code block
        if line.strip().startswith('```'):
            lang = line.strip()[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            blocks.append({'type': 'code', 'lang': lang, 'content': '\n'.join(code_lines)})
            continue
        # Heading
        h = re.match(r'^(#{1,6})\s+(.+)$', line)
        if h:
            blocks.append({'type': 'heading', 'level': len(h.group(1)), 'content': h.group(2)})
            i += 1
            continue
        # Table
        if i + 1 < len(lines) and re.match(r'^\|[\s\-:|]+\|$', lines[i+1].strip()):
            header_cells = [c.strip() for c in line.strip().split('|')[1:-1]]
            aligns = []
            for s in re.findall(r':?-{3,}:?', lines[i+1]):
                if s.startswith(':') and s.endswith(':'):
                    aligns.append('center')
                elif s.endswith(':'):
                    aligns.append('right')
                else:
                    aligns.append('left')
            rows = []
            i += 2
            while i < len(lines) and lines[i].strip().startswith('|'):
                cells = [c.strip() for c in lines[i].strip().split('|')[1:-1]]
                rows.append(cells)
                i += 1
            blocks.append({'type': 'table', 'headers': header_cells, 'aligns': aligns, 'rows': rows})
            continue
        # Blockquote
        if line.strip().startswith('>'):
            quote_lines = []
            while i < len(lines) and (lines[i].strip().startswith('>') or
                  (lines[i].strip() == '' and i+1 < len(lines) and lines[i+1].strip().startswith('>'))):
                if lines[i].strip().startswith('>'):
                    quote_lines.append(re.sub(r'^>\s?', '', lines[i]))
                i += 1
            blocks.append({'type': 'blockquote', 'content': '\n'.join(quote_lines)})
            continue
        # Unordered list
        if re.match(r'^[\s]*[-*+]\s+', line):
            items = []
            while i < len(lines) and re.match(r'^[\s]*[-*+]\s+', lines[i]):
                items.append(re.sub(r'^[\s]*[-*+]\s+', '', lines[i]))
                i += 1
            blocks.append({'type': 'ul', 'items': items})
            continue
        # Ordered list
        if re.match(r'^[\s]*\d+[.)]\s+', line):
            items = []
            while i < len(lines) and re.match(r'^[\s]*\d+[.)]\s+', lines[i]):
                items.append(re.sub(r'^[\s]*\d+[.)]\s+', '', lines[i]))
                i += 1
            blocks.append({'type': 'ol', 'items': items})
            continue
        # Paragraph
        para_lines = []
        while i < len(lines) and lines[i].strip() != '' and \
              not re.match(r'^(#{1,6}\s|```|\||>|[-*+]\s|\d+[.)]\s|^[-*_]{3,})', lines[i]):
            para_lines.append(lines[i])
            i += 1
        blocks.append({'type': 'para', 'content': ' '.join(para_lines)})
    return blocks

# ── inline formatting parser ─────────────────────────────────────────
def parse_inline(text):
    """Parse **bold**, *italic*, `code`, $math$, [links]()"""
    result = []
    remaining = text
    while remaining:
        # Display math $$...$$
        m = re.match(r'(.*?)\$\$(.+?)\$\$', remaining, re.DOTALL)
        if m:
            if m.group(1):
                result.append(('text', m.group(1)))
            result.append(('display_math', m.group(2)))
            remaining = remaining[m.end():]
            continue
        # Inline math $...$ (lookbehind for non-escaped $)
        m = re.match(r'(.*?)(?:(?<!\\)\$)(.+?)(?:(?<!\\)\$)', remaining)
        if m and m.group(2):
            if m.group(1):
                result.append(('text', m.group(1)))
            result.append(('math', m.group(2)))
            remaining = remaining[m.end():]
            continue
        # Link [text](url)
        m = re.match(r'(.*?)\[(.+?)\]\((.+?)\)', remaining)
        if m:
            if m.group(1):
                result.append(('text', m.group(1)))
            result.append(('link', m.group(2), m.group(3)))
            remaining = remaining[m.end():]
            continue
        # Bold **...**
        m = re.match(r'(.*?)\*\*(.+?)\*\*', remaining)
        if m:
            if m.group(1):
                result.append(('text', m.group(1)))
            result.append(('bold', m.group(2)))
            remaining = remaining[m.end():]
            continue
        # Italic *...*
        m = re.match(r'(.*?)\*(.+?)\*', remaining)
        if m:
            if m.group(1):
                result.append(('text', m.group(1)))
            result.append(('italic', m.group(2)))
            remaining = remaining[m.end():]
            continue
        # Inline code `...`
        m = re.match(r'(.*?)`(.+?)`', remaining)
        if m:
            if m.group(1):
                result.append(('text', m.group(1)))
            result.append(('code', m.group(2)))
            remaining = remaining[m.end():]
            continue
        # Plain text
        result.append(('text', remaining))
        break
    return result

# ── LaTeX helpers ────────────────────────────────────────────────────
def tex_escape(text):
    for char, repl in [('&', r'\&'), ('%', r'\%'), ('#', r'\#'),
                       ('{', r'\{'), ('}', r'\}'), ('~', r'\textasciitilde{}'),
                       ('$', r'\$'), ('_', r'\_'), ('^', r'\^{}')]:
        text = text.replace(char, repl)
    return text

def inline_to_tex(tokens):
    parts = []
    for tok in tokens:
        kind = tok[0]
        if kind == 'text':
            parts.append(tex_escape(tok[1]))
        elif kind == 'bold':
            parts.append(r'\textbf{' + tex_escape(tok[1]) + '}')
        elif kind == 'italic':
            parts.append(r'\textit{' + tex_escape(tok[1]) + '}')
        elif kind == 'code':
            parts.append(r'\texttt{' + tex_escape(tok[1]) + '}')
        elif kind == 'math':
            parts.append('$' + tok[1] + '$')
        elif kind == 'display_math':
            parts.append(r'\[' + tok[1] + r'\]')
        elif kind == 'link':
            url_esc = tok[2].replace('%', r'\%').replace('#', r'\#')
            parts.append(r'\href{' + url_esc + '}{' + tex_escape(tok[1]) + '}')
    return ''.join(parts)

# ── DOCX generation ──────────────────────────────────────────────────
def generate_docx(blocks, outpath):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    for block in blocks:
        t = block['type']
        if t == 'heading':
            h = doc.add_heading(level=min(block['level'], 3))
            run = h.add_run(block['content'])
            run.font.name = 'Calibri'
        elif t == 'hr':
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '999999')
            pBdr.append(bottom)
            pPr.append(pBdr)
        elif t == 'para':
            p = doc.add_paragraph()
            add_runs_docx(p, parse_inline(block['content']))
        elif t == 'blockquote':
            for part in block['content'].split('\n\n'):
                if part.strip():
                    pq = doc.add_paragraph()
                    pq.style = doc.styles['Quote']
                    add_runs_docx(pq, parse_inline(part.strip()))
        elif t == 'code':
            p = doc.add_paragraph()
            p.style = doc.styles['No Spacing']
            run = p.add_run(block['content'])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), 'F0F0F0')
            shd.set(qn('w:val'), 'clear')
            pPr.append(shd)
        elif t == 'ul':
            for item in block['items']:
                p = doc.add_paragraph(style='List Bullet')
                add_runs_docx(p, parse_inline(item))
        elif t == 'ol':
            for item in block['items']:
                p = doc.add_paragraph(style='List Number')
                add_runs_docx(p, parse_inline(item))
        elif t == 'table':
            nrows = len(block['rows']) + 1
            ncols = len(block['headers'])
            table = doc.add_table(rows=nrows, cols=ncols)
            table.style = 'Light Grid Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ci, ct in enumerate(block['headers']):
                cell = table.rows[0].cells[ci]
                cell.text = ''
                r = cell.paragraphs[0].add_run(ct)
                r.bold = True
                r.font.size = Pt(10)
            for ri, row in enumerate(block['rows']):
                for ci, ct in enumerate(row):
                    cell = table.rows[ri+1].cells[ci]
                    cell.text = ''
                    add_runs_docx(cell.paragraphs[0], parse_inline(ct))
    doc.save(outpath)
    print(f"DOCX saved: {outpath}")

def add_runs_docx(paragraph, tokens):
    from docx.shared import Pt
    for tok in tokens:
        kind = tok[0]
        if kind == 'text':
            paragraph.add_run(tok[1])
        elif kind == 'bold':
            r = paragraph.add_run(tok[1]); r.bold = True
        elif kind == 'italic':
            r = paragraph.add_run(tok[1]); r.italic = True
        elif kind == 'code':
            r = paragraph.add_run(tok[1]); r.font.name = 'Consolas'; r.font.size = Pt(9)
        elif kind in ('math', 'display_math'):
            paragraph.add_run('$' + tok[1] + '$')
        elif kind == 'link':
            r = paragraph.add_run(tok[1])
            r.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            r.underline = True

# ── LaTeX / PDF generation ───────────────────────────────────────────
def generate_latex(blocks, outpath):
    lines = []
    lines.append(r'\documentclass[11pt,a4paper]{article}')
    lines.append(r'\usepackage[utf8]{inputenc}')
    lines.append(r'\usepackage[T1]{fontenc}')
    lines.append(r'\usepackage{amsmath,amssymb}')
    lines.append(r'\usepackage[margin=2.5cm]{geometry}')
    lines.append(r'\usepackage{booktabs}')
    lines.append(r'\usepackage{longtable}')
    lines.append(r'\usepackage{array}')
    lines.append(r'\usepackage{enumitem}')
    lines.append(r'\usepackage{hyperref}')
    lines.append(r'\usepackage{xcolor}')
    lines.append(r'\usepackage{fancyhdr}')
    lines.append(r'\usepackage{parskip}')
    lines.append(r'\hypersetup{colorlinks=true,linkcolor=blue,urlcolor=blue,citecolor=blue}')
    lines.append(r'\setlength{\parindent}{0pt}')
    lines.append(r'\setlength{\parskip}{6pt}')
    lines.append(r'\pagestyle{fancy}')
    lines.append(r'\fancyhf{}')
    lines.append(r'\fancyhead[L]{\small QWAV $\times$ apoth3osis -- Meeting Agenda}')
    lines.append(r'\fancyhead[R]{\small May 2026}')
    lines.append(r'\fancyfoot[C]{\thepage}')
    lines.append(r'\begin{document}')
    lines.append('')

    for block in blocks:
        t = block['type']
        if t == 'heading':
            cmd = {1: 'section*', 2: 'subsection*', 3: 'subsubsection*',
                   4: 'paragraph*', 5: 'paragraph*', 6: 'paragraph*'}[block['level']]
            text = inline_to_tex(parse_inline(block['content']))
            # Use raw string concatenation to avoid escaping issues
            line = '\\' + cmd + '{' + text + '}'
            lines.append(line)
            lines.append('')
        elif t == 'hr':
            lines.append(r'\bigskip')
            lines.append(r'\noindent\rule{\textwidth}{0.4pt}')
            lines.append(r'\bigskip')
        elif t == 'para':
            text = inline_to_tex(parse_inline(block['content']))
            lines.append(text)
            lines.append('')
        elif t == 'blockquote':
            lines.append(r'\begin{quote}')
            for part in block['content'].split('\n\n'):
                if part.strip():
                    text = inline_to_tex(parse_inline(part.strip()))
                    lines.append(text + r'\par')
            lines.append(r'\end{quote}')
            lines.append('')
        elif t == 'code':
            lines.append(r'\begin{verbatim}')
            lines.append(block['content'])
            lines.append(r'\end{verbatim}')
            lines.append('')
        elif t == 'ul':
            lines.append(r'\begin{itemize}[nosep]')
            for item in block['items']:
                text = inline_to_tex(parse_inline(item))
                lines.append(r'\item ' + text)
            lines.append(r'\end{itemize}')
            lines.append('')
        elif t == 'ol':
            lines.append(r'\begin{enumerate}[nosep]')
            for item in block['items']:
                text = inline_to_tex(parse_inline(item))
                lines.append(r'\item ' + text)
            lines.append(r'\end{enumerate}')
            lines.append('')
        elif t == 'table':
            col_spec = '|' + '|'.join(
                'l' if a == 'left' else ('c' if a == 'center' else 'r')
                for a in block['aligns']
            ) + '|'
            lines.append(r'\begin{center}')
            lines.append(r'\begin{longtable}{' + col_spec + '}')
            lines.append(r'\hline')
            hcells = [r'\textbf{' + inline_to_tex(parse_inline(h)) + '}' for h in block['headers']]
            lines.append(' & '.join(hcells) + r' \\ \hline')
            lines.append(r'\endhead')
            for row in block['rows']:
                cells = [inline_to_tex(parse_inline(c)) for c in row]
                lines.append(' & '.join(cells) + r' \\ \hline')
            lines.append(r'\end{longtable}')
            lines.append(r'\end{center}')
            lines.append('')

    lines.append(r'\end{document}')

    with open(outpath, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    print(f"LaTeX saved: {outpath}")

def compile_pdf(tex_path, out_dir):
    pdflatex = os.path.join(TEXLIVE_BIN, 'pdflatex.exe')
    orig_cwd = os.getcwd()
    os.chdir(out_dir)
    for _ in range(2):
        result = subprocess.run(
            [pdflatex, '-interaction=nonstopmode', '-halt-on-error', tex_path],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode != 0:
            err_lines = [l for l in (result.stdout + result.stderr).split('\n') if l.startswith('!')]
            if err_lines:
                print(f"LaTeX warnings (non-fatal with nonstopmode):")
                for el in err_lines[:5]:
                    print(f"  {el}")
    os.chdir(orig_cwd)
    pdf_path = os.path.join(out_dir, os.path.splitext(os.path.basename(tex_path))[0] + '.pdf')
    if os.path.exists(pdf_path):
        print(f"PDF compiled: {pdf_path}")
        return pdf_path
    else:
        print("PDF compilation FAILED.")
        return None

# ── main ─────────────────────────────────────────────────────────────
def main():
    text = read_md(SRC)
    blocks = parse_md(text)
    print(f"Parsed {len(blocks)} blocks")

    docx_path = BASE + ".docx"
    generate_docx(blocks, docx_path)

    tex_path = BASE + ".tex"
    generate_latex(blocks, tex_path)
    pdf_path = compile_pdf(tex_path, OUT_DIR)

    if pdf_path:
        print(f"\nDone! DOCX: {docx_path}  |  PDF: {pdf_path}")
    else:
        print(f"\nDOCX created at {docx_path}. PDF failed.")

if __name__ == '__main__':
    main()
